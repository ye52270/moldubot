from __future__ import annotations

import os
import uuid
from pathlib import Path

from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.core.logging_config import get_logger

ROOT_DIR = Path(__file__).resolve().parents[2]
REPORT_FILES_DIR = ROOT_DIR / "data" / "reports" / "docx"
REPORT_HTML_DIR = ROOT_DIR / "data" / "reports" / "html"
logger = get_logger(__name__)


def _safe_file_stem(title: str) -> str:
    """
    제목 문자열을 파일명으로 사용할 수 있는 안전한 형태로 변환한다.

    Args:
        title: 보고서 제목

    Returns:
        파일명 stem 문자열
    """
    normalized = "".join(ch for ch in str(title or "") if ch.isalnum() or ch in " -_")
    stem = normalized.strip().replace(" ", "_")
    return (stem or "report")[:40]


def _shade_cell(cell, fill: str) -> None:
    """
    DOCX 테이블 셀 배경색을 지정한다.

    Args:
        cell: python-docx 셀 객체
        fill: 6자리 HEX 색상
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _shade_paragraph(paragraph, fill: str) -> None:
    """
    DOCX 문단 배경색을 지정한다.

    Args:
        paragraph: python-docx 문단 객체
        fill: 6자리 HEX 색상
    """
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def _add_table(doc: Document, table_tag: Tag) -> None:
    """
    HTML table 태그를 DOCX 표로 변환한다.

    Args:
        doc: Word 문서 객체
        table_tag: BeautifulSoup table 태그
    """
    rows = table_tag.find_all("tr")
    if not rows:
        return
    col_count = max(len(row.find_all(["td", "th"])) for row in rows)
    if col_count <= 0:
        return
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for col_index, cell in enumerate(cells[:col_count]):
            target = table.rows[row_index].cells[col_index]
            target.text = cell.get_text(strip=True)
            if cell.name == "th" or row_index == 0:
                _shade_cell(target, "1F4E79")
                for paragraph in target.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True


def _process_nodes(doc: Document, element: Tag) -> None:
    """
    HTML 노드를 순회하며 DOCX 콘텐츠를 생성한다.

    Args:
        doc: Word 문서 객체
        element: BeautifulSoup 부모 노드
    """
    for child in element.children:
        if not isinstance(child, Tag):
            continue
        tag = child.name.lower()
        if tag == "h1":
            paragraph = doc.add_heading(child.get_text(strip=True), level=0)
            if paragraph.runs:
                paragraph.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif tag == "h2":
            paragraph = doc.add_heading(child.get_text(strip=True), level=1)
            if paragraph.runs:
                paragraph.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        elif tag == "h3":
            doc.add_heading(child.get_text(strip=True), level=2)
        elif tag == "p":
            doc.add_paragraph(child.get_text(strip=True))
        elif tag in ("ul", "ol"):
            for li_tag in child.find_all("li", recursive=False):
                doc.add_paragraph(li_tag.get_text(strip=True), style="List Bullet")
        elif tag == "table":
            _add_table(doc=doc, table_tag=child)
        elif tag == "div":
            classes = child.get("class", [])
            if any(name in classes for name in ("warning", "alert", "risk")):
                paragraph = doc.add_paragraph(f"⚠ {child.get_text(strip=True)}")
                _shade_paragraph(paragraph, "FFF3CD")
            else:
                _process_nodes(doc=doc, element=child)
        elif tag == "hr":
            doc.add_paragraph("─" * 50)
        else:
            _process_nodes(doc=doc, element=child)


def _configure_document_layout(document: Document, layout: str) -> None:
    """
    보고서 레이아웃 타입에 따라 문서 페이지 설정을 적용한다.

    Args:
        document: python-docx 문서 객체
        layout: `portrait` 또는 `landscape_wide`
    """
    section = document.sections[0]
    normalized = str(layout or "portrait").strip().lower()
    if normalized == "landscape_wide":
        section.page_width = Cm(33.867)
        section.page_height = Cm(19.05)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        return
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


async def convert_html_to_docx(html: str, title: str, layout: str = "portrait") -> str:
    """
    HTML 보고서를 DOCX로 변환하고 다운로드 경로를 반환한다.

    Args:
        html: 보고서 HTML
        title: 파일명에 사용할 제목
        layout: 페이지 레이아웃(`portrait` 또는 `landscape_wide`)

    Returns:
        다운로드 URL 경로(`/report/download/{filename}`)
    """
    document = Document()
    _configure_document_layout(document=document, layout=layout)
    document.styles["Normal"].font.name = "맑은 고딕"
    document.styles["Normal"].font.size = Pt(10.5)

    soup = BeautifulSoup(str(html or ""), "lxml")
    root = soup.body if isinstance(soup.body, Tag) else soup
    _process_nodes(doc=document, element=root)

    REPORT_FILES_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_HTML_DIR.mkdir(parents=True, exist_ok=True)
    file_stem = f"{uuid.uuid4().hex[:8]}_{_safe_file_stem(title)}"
    filename = f"{file_stem}.docx"
    file_path = REPORT_FILES_DIR / filename
    document.save(str(file_path))
    html_path = REPORT_HTML_DIR / f"{file_stem}.html"
    html_path.write_text(str(html or ""), encoding="utf-8")
    logger.info("report_docx 변환 완료: file=%s", file_path.name)
    return f"/report/download/{filename}"


def resolve_report_file_path(filename: str) -> Path:
    """
    다운로드 요청 파일명을 실제 경로로 변환한다.

    Args:
        filename: 요청 파일명

    Returns:
        docx 파일 절대 경로
    """
    safe_name = os.path.basename(str(filename or "").strip())
    return REPORT_FILES_DIR / safe_name


def resolve_report_html_path(filename: str) -> Path:
    """
    다운로드 파일명을 HTML 미리보기 경로로 변환한다.

    Args:
        filename: docx 파일명 또는 stem

    Returns:
        html 파일 절대 경로
    """
    safe_name = os.path.basename(str(filename or "").strip())
    html_name = f"{Path(safe_name).stem}.html"
    return REPORT_HTML_DIR / html_name
